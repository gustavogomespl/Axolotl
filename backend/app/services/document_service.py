import os
import tempfile
import uuid

from fastapi import UploadFile
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.vector_store.client import VectorStoreManager


class DocumentService:
    """Pipeline for document ingestion:
    1. Upload (PDF, TXT, MD, DOCX)
    2. Parse content
    3. Chunk with RecursiveCharacterTextSplitter
    4. Embed and index in ChromaDB
    5. Track metadata
    """

    def __init__(self, vector_store: VectorStoreManager | None = None):
        self.vector_store = vector_store or VectorStoreManager()

    async def ingest(
        self,
        file: UploadFile,
        collection: str,
        metadata: dict | None = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> dict:
        """Ingest a file into the vector store."""
        metadata = metadata or {}
        doc_id = str(uuid.uuid4())

        # Save file temporarily
        suffix = os.path.splitext(file.filename or "")[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name

        try:
            # Parse based on file type
            text = await self._parse_file(tmp_path, suffix)

            # Chunk
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            chunks = splitter.split_text(text)

            # Create LangChain documents
            documents = [
                LCDocument(
                    page_content=chunk,
                    metadata={
                        "document_id": doc_id,
                        "filename": file.filename,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        **metadata,
                    },
                )
                for i, chunk in enumerate(chunks)
            ]

            # Add to vector store
            self.vector_store.add_documents(collection, documents)

            return {
                "id": doc_id,
                "filename": file.filename,
                "collection": collection,
                "chunk_count": len(chunks),
                "status": "ready",
            }
        finally:
            os.unlink(tmp_path)

    async def _parse_file(self, file_path: str, suffix: str) -> str:
        """Parse file content based on type."""
        suffix = suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix in (".txt", ".md"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif suffix == ".docx":
            return self._parse_docx(file_path)
        else:
            # Fallback: try reading as text
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF using pdfplumber."""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            # Fallback to basic PDF reading
            from langchain_community.document_loaders import PyPDFLoader

            loader = PyPDFLoader(file_path)
            docs = loader.load()
            return "\n\n".join(doc.page_content for doc in docs)

    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX using python-docx."""
        import docx

        doc = docx.Document(file_path)
        return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())

    async def delete(self, collection: str, document_id: str) -> None:
        """Delete a document from the vector store by its ID."""
        collection_obj = self.vector_store.get_or_create_collection(collection)
        collection_obj.delete(where={"document_id": document_id})
